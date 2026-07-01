"""Resume adapter (unstructured -> PDF / DOCX prose).

Two stages: (1) pull plain text out of the file -- pdfplumber for PDF, python-docx
for DOCX; (2) heuristically extract fields from that text. Contact details, links
and skills reuse the recruiter-notes patterns; experience and education are parsed
from their labelled sections using common, documented layouts.

Resume data is openly lower-trust (method=regex), so it never out-ranks the ATS.
Anything that doesn't match a known pattern is skipped, not guessed -- a garbled or
image-only PDF simply yields fewer fields, never a crash.

When GROQ_API_KEY is configured, an optional LLM extractor (see transformer.extract)
runs on the same text and emits a second, `llm`-tagged record for what the
heuristics miss (prose skills, messy experience/education). Its output is verified
against the source text, so it augments -- never fabricates -- and the two records
are merged by the normal engine. Without a key, only the heuristic record is used
and runs stay fully deterministic.
"""

from __future__ import annotations

import os
import re
from typing import List, Optional

from ..model import SourceRecord, SOURCE_RESUME, METHOD_REGEX, METHOD_LLM
from .notes_source import (
    _EMAIL_RE, _PHONE_RE, _GITHUB_RE, _LINKEDIN_RE, _YEARS_RE,
    _find_skills, _guess_name,
)

_PORTFOLIO_RE = re.compile(r"(?:portfolio|website)\s*[:\-]?\s*(https?://\S+)", re.I)
# "Staff Engineer at Acme Corp (2021-03 to present)"
_EXP_RE = re.compile(r"^(?P<title>.+?)\s+at\s+(?P<company>.+?)\s*\((?P<span>[^)]+)\)\s*$", re.I)
# "B.S. in Computer Science, MIT, 2017"
_EDU_RE = re.compile(r"^(?P<deg>.+?),\s*(?P<inst>.+?),\s*(?P<year>\d{4})\s*$")

# Known section headings -> used as section boundaries (real resumes rarely leave
# blank lines between sections, and headings are often ALL-CAPS).
_HEADINGS = {
    "education", "experience", "work experience", "professional experience",
    "employment", "skills", "technical skills", "projects", "certifications",
    "summary", "objective", "awards", "publications", "languages", "interests",
    "achievements", "publications & achievements", "contact", "profile",
}

# A degree token (B.E., B.Tech, M.S., PhD, Bachelor, ...).
_DEGREE_RE = re.compile(
    r"\b(?:Ph\.?D|B\.?E|B\.?Tech|B\.?Sc|B\.?S|B\.?A|M\.?E|M\.?Tech|M\.?Sc|M\.?S|"
    r"M\.?A|M\.?B\.?A|Bachelor|Master|Diploma)\b\.?", re.I,
)

# A single date token, then a start–end range (handles "Dec 2025 - May 2026",
# "2021-03 to present", "2019").
_DTOK = r"(?:[A-Za-z]{3,9}\.?\s+\d{4}|\d{4}[-/]\d{1,2}|\d{4})"
_DATERANGE_RE = re.compile(
    r"(%s)\s*(?:-|–|—|to|until)\s*(%s|present|current|now|ongoing)" % (_DTOK, _DTOK), re.I,
)
_YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")


# --------------------------------------------------------------------------- #
# Text extraction
# --------------------------------------------------------------------------- #
def _clean(text: str) -> str:
    # PDF fonts often emit unicode dashes/bullets (or the U+FFFD replacement char)
    # that terminals mangle and our patterns miss; fold them to ASCII "-" so date
    # ranges, separators and bullet markers stay parseable.
    for ch in ("�", "–", "—", "―", "•", "·", "●", "‣"):
        text = text.replace(ch, "-")
    return text


def extract_text(path: str) -> str:
    low = path.lower()
    try:
        if low.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                return _clean("\n".join((page.extract_text() or "") for page in pdf.pages))
        if low.endswith(".docx"):
            import docx
            doc = docx.Document(path)
            return _clean("\n".join(p.text for p in doc.paragraphs))
    except Exception:
        # Corrupt / image-only / unreadable file -> no text, no crash.
        return ""
    return ""


def extract_hyperlinks(path: str) -> List[str]:
    """Return embedded hyperlink URLs. Résumés commonly show 'LinkedIn'/'GitHub' as
    clickable text with the real URL only in the PDF/DOCX link annotation, so those
    are invisible to text extraction -- we read them from the annotations directly."""
    low, urls = path.lower(), []
    try:
        if low.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    for h in (page.hyperlinks or []):
                        if h.get("uri"):
                            urls.append(h["uri"])
                    for a in (page.annots or []):
                        uri = a.get("uri") or (a.get("data") or {}).get("A", {}).get("URI")
                        if uri:
                            urls.append(uri.decode("latin-1", "ignore") if isinstance(uri, bytes) else str(uri))
        elif low.endswith(".docx"):
            import docx
            doc = docx.Document(path)
            for rel in doc.part.rels.values():
                if "hyperlink" in rel.reltype.lower():
                    urls.append(rel.target_ref)
    except Exception:
        return []
    seen = []
    for u in urls:
        u = str(u).strip()
        if u and u not in seen:
            seen.append(u)
    return seen


def _links_from_urls(urls: List[str]) -> dict:
    """Classify hyperlink URLs into github / linkedin (high precision only)."""
    out = {}
    for u in urls:
        ul = u.lower().rstrip("/")
        if "github.com/" in ul and "github" not in out:
            handle = u.rstrip("/").split("github.com/")[-1].split("/")[0].split("?")[0]
            if handle and handle.lower() not in ("about", "login", "features"):
                out["github"] = handle
        elif "linkedin.com/in/" in ul and "linkedin" not in out:
            out["linkedin"] = u.split("?")[0]
    return out


# --------------------------------------------------------------------------- #
# Section-aware field extraction
# --------------------------------------------------------------------------- #
def _is_heading(line: str) -> bool:
    s = line.strip().lower().rstrip(":")
    if s in _HEADINGS:
        return True
    # ALL-CAPS short line (e.g. "TECHNICAL SKILLS", "PUBLICATIONS & ACHIEVEMENTS").
    raw = line.strip()
    return bool(raw) and raw.upper() == raw and len(raw.split()) <= 3 and len(raw) <= 30 and any(c.isalpha() for c in raw)


def _section(lines: List[str], *names: str) -> List[str]:
    """Lines under any of ``names`` until the next known heading."""
    wanted = {n.lower() for n in names}
    out, capturing = [], False
    for line in lines:
        key = line.strip().lower().rstrip(":")
        if not capturing:
            if key in wanted:
                capturing = True
            continue
        if _is_heading(line):
            break
        if line.strip():
            out.append(line.strip())
    return out


def _is_bullet(line: str) -> bool:
    return line[:1] in "-•*·"


def _parse_span(span: str):
    parts = re.split(r"\s+(?:to|until|–|—)\s+", span, maxsplit=1)
    start = parts[0].strip() if parts else None
    end = parts[1].strip() if len(parts) > 1 else None
    return start, end


def _extract_skills(lines: List[str], full_text: str) -> List[str]:
    """Open-vocabulary: take everything listed in a Skills section (known or not),
    then union with the keyword scan for skills mentioned elsewhere in prose."""
    found: List[str] = []
    for line in _section(lines, "Technical Skills", "Skills"):
        # Drop a leading category label ("Languages:", "AI / ML:").
        if ":" in line:
            label, rest = line.split(":", 1)
            if len(label) <= 25 and "," not in label:
                line = rest
        line = re.sub(r"\([^)]*\)", "", line)  # drop parenthetical detail
        for tok in re.split(r"[,;|•]", line):
            tok = tok.strip(" .-")
            if tok and len(tok) <= 40 and re.search(r"[A-Za-z]", tok):
                found.append(tok)
    # union with the closed-vocabulary scan (catches skills named only in prose)
    found += _find_skills(full_text)
    seen, out = set(), []
    for s in found:
        k = s.lower()
        if k not in seen:
            seen.add(k); out.append(s)
    return out


def _parse_experience(lines: List[str]) -> List[dict]:
    exp = []
    for line in _section(lines, "Experience", "Work Experience", "Professional Experience"):
        if _is_bullet(line):
            continue
        m = _EXP_RE.match(line)
        if m:  # strict "Title at Company (dates)"
            start, end = _parse_span(m.group("span"))
            exp.append({"company": m.group("company").strip(), "title": m.group("title").strip(),
                        "start": start, "end": end, "summary": None})
            continue
        # Flexible: pipe-delimited with a trailing date range.
        dm = _DATERANGE_RE.search(line)
        if not dm:
            continue
        head = line[:dm.start()].strip(" |-–—")
        segs = [s.strip() for s in re.split(r"\s*\|\s*", head) if s.strip()]
        if not segs:
            continue
        exp.append({"title": segs[0], "company": segs[1] if len(segs) > 1 else None,
                    "start": dm.group(1), "end": dm.group(2), "summary": None})
    return exp


def _parse_education(lines: List[str]) -> List[dict]:
    section = _section(lines, "Education")
    if not section:
        return []
    # A year may sit on its own line ("Expected: October 2027").
    section_year = None
    ym = _YEAR_RE.search(" ".join(section))
    if ym:
        section_year = int(ym.group())

    edu = []
    for line in section:
        if _is_bullet(line):
            continue
        m = _EDU_RE.match(line)
        if m:  # strict "Degree in Field, Institution, Year"
            deg = m.group("deg").strip()
            degree, field = (deg.split(" in ", 1) + [None])[:2] if " in " in deg else (deg, None)
            edu.append({"institution": m.group("inst").strip(), "degree": degree.strip(),
                        "field": field.strip() if field else None, "end_year": int(m.group("year"))})
            continue
        # Flexible: a line with a degree token and/or pipe/comma structure.
        if not (_DEGREE_RE.search(line) or "|" in line):
            continue
        segs = [s.strip() for s in re.split(r"\s*\|\s*", line) if s.strip()]
        deg_field = segs[0] if segs else line
        dm = _DEGREE_RE.search(deg_field)
        degree = dm.group(0).strip() if dm else None
        field = deg_field[dm.end():].strip(" -–—:,") if dm else None
        institution = None
        for s in segs[1:]:
            if re.search(r"(University|Institute|College|School|Polytechnic|Academy)", s, re.I):
                institution = s.split(",")[0].strip()
                break
        if institution is None and len(segs) > 1:
            institution = segs[1].split(",")[0].strip()
        yr = _YEAR_RE.search(line)
        edu.append({"institution": institution, "degree": degree,
                    "field": field or None, "end_year": int(yr.group()) if yr else section_year})
    return edu


def parse_text(text: str) -> dict:
    raw = {}
    if not text.strip():
        return raw
    lines = text.splitlines()

    emails = _EMAIL_RE.findall(text)
    if emails:
        raw["emails"] = emails
    phones = [p.strip() for p in _PHONE_RE.findall(text)]
    if phones:
        raw["phones"] = phones

    links = {}
    gh = _GITHUB_RE.search(text)
    if gh and (gh.group(1) or gh.group(2)):
        links["github"] = gh.group(1) or gh.group(2)
    li = _LINKEDIN_RE.search(text)
    if li:
        links["linkedin"] = "https://www.linkedin.com/in/" + li.group(1)
    pf = _PORTFOLIO_RE.search(text)
    if pf:
        links["portfolio"] = pf.group(1)
    if links:
        raw["links"] = links

    name = _guess_name(text)
    if name:
        raw["full_name"] = name

    ym = _YEARS_RE.search(text)
    if ym:
        try:
            raw["years_experience"] = float(ym.group(1))
        except ValueError:
            pass

    skills = _extract_skills(lines, text)
    if skills:
        raw["skills"] = skills

    experience = _parse_experience(lines)
    if experience:
        raw["experience"] = experience

    education = _parse_education(lines)
    if education:
        raw["education"] = education

    return raw


def _llm_extract(text: str) -> Optional[dict]:
    """Grounded LLM extraction; returns None if disabled or on any failure."""
    try:
        from ..extract import llm_resume
        return llm_resume.extract(text)
    except Exception:
        return None


def load(path: str, use_llm: Optional[bool] = None) -> List[SourceRecord]:
    """Load a single resume file or a directory of resumes (.pdf / .docx).

    Always emits a heuristic (method=regex) record. When ``use_llm`` is true -- or
    left None and GROQ_API_KEY is set -- it also emits a verified LLM (method=llm)
    record for the same file, and the engine merges the two.
    """
    if use_llm is None:
        use_llm = bool(os.environ.get("GROQ_API_KEY"))

    paths: List[str] = []
    if os.path.isdir(path):
        for name in sorted(os.listdir(path)):
            if name.lower().endswith((".pdf", ".docx")):
                paths.append(os.path.join(path, name))
    else:
        paths.append(path)

    records: List[SourceRecord] = []
    for p in paths:
        text = extract_text(p)
        raw = parse_text(text)
        # Merge in links read from embedded hyperlink annotations (github/linkedin
        # are often clickable text with no visible URL). Hyperlinks are the actual
        # target, so they win over any text-guessed handle.
        hyper = _links_from_urls(extract_hyperlinks(p))
        if hyper:
            links = raw.get("links", {})
            links.update(hyper)
            raw["links"] = links
        if raw:
            records.append(SourceRecord(SOURCE_RESUME, raw, {k: METHOD_REGEX for k in raw}))
        if use_llm:
            llm_raw = _llm_extract(text)
            if llm_raw:
                records.append(SourceRecord(SOURCE_RESUME, llm_raw, {k: METHOD_LLM for k in llm_raw}))
    return records
